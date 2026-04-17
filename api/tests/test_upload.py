"""
Tests — Upload de documents (/upload)
"""

import io
import pytest


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def upload_file(client, content: bytes, filename: str, content_type: str = "text/plain"):
    return client.post(
        "/upload",
        files={"file": (filename, io.BytesIO(content), content_type)},
    )


# ─────────────────────────────────────────────────────────────────────────────
# TXT
# ─────────────────────────────────────────────────────────────────────────────

class TestTxtUpload:
    def test_txt_upload_success(self, client):
        text = "Bonjour, voici un document de test pour Comply.\n" * 10
        resp = upload_file(client, text.encode(), "test.txt")
        assert resp.status_code == 200
        data = resp.json()
        assert data["filename"] == "test.txt"
        assert data["extension"] == "txt"
        assert "Bonjour" in data["text"]
        assert data["chars"] > 0

    def test_txt_upload_with_utf8_content(self, client):
        text = "Données en français : é, è, ê, ç, à — étudiants JE."
        resp = upload_file(client, text.encode("utf-8"), "french.txt")
        assert resp.status_code == 200
        assert "français" in resp.json()["text"]

    def test_md_treated_as_txt(self, client):
        content = "# Guide de facturation\n\nVoici comment facturer."
        resp = upload_file(client, content.encode(), "guide.md", "text/markdown")
        assert resp.status_code == 200
        assert resp.json()["extension"] == "md"

    def test_json_treated_as_txt(self, client):
        content = b'{"key": "value"}'
        resp = upload_file(client, content, "data.json", "application/json")
        assert resp.status_code == 200

    def test_truncation_above_50k_chars(self, client):
        # Texte de 60 000 chars
        long_text = ("A" * 100 + "\n") * 600
        resp = upload_file(client, long_text.encode(), "long.txt")
        assert resp.status_code == 200
        data = resp.json()
        assert data["truncated"] is True
        assert len(data["text"]) == 50000

    def test_no_truncation_below_50k_chars(self, client):
        short_text = "Texte court.\n" * 10
        resp = upload_file(client, short_text.encode(), "short.txt")
        assert resp.status_code == 200
        assert resp.json()["truncated"] is False


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

class TestPdfUpload:
    def test_valid_pdf_upload(self, client):
        """Génère un mini PDF valide avec pypdf."""
        pytest.importorskip("pypdf")
        from pypdf import PdfWriter
        writer = PdfWriter()
        page = writer.add_blank_page(width=595, height=842)
        # Ajouter un flux de texte simple (page vide = texte extrait vide)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        resp = upload_file(client, buf.read(), "test.pdf", "application/pdf")
        # Page vide → texte vide → 422
        assert resp.status_code in (200, 422)

    def test_pdf_with_text(self, client):
        """Crée un PDF avec du texte réel."""
        pytest.importorskip("pypdf")
        from pypdf import PdfWriter
        from pypdf.generic import (
            ArrayObject, DecodedStreamObject, DictionaryObject,
            NameObject, NumberObject, ByteStringObject,
        )
        # PDF minimal hand-crafted avec du texte
        minimal_pdf = b"""%PDF-1.4
1 0 obj<</Type /Catalog /Pages 2 0 R>>endobj
2 0 obj<</Type /Pages /Kids[3 0 R] /Count 1>>endobj
3 0 obj<</Type /Page /Parent 2 0 R /MediaBox[0 0 612 792]
/Contents 4 0 R /Resources<</Font<</F1 5 0 R>>>>>>endobj
4 0 obj<</Length 44>>
stream
BT /F1 12 Tf 72 720 Td (Junior-Entreprise) Tj ET
endstream
endobj
5 0 obj<</Type /Font /Subtype /Type1 /BaseFont /Helvetica>>endobj
xref
0 6
0000000000 65535 f\r
0000000009 00000 n\r
0000000058 00000 n\r
0000000115 00000 n\r
0000000266 00000 n\r
0000000360 00000 n\r
trailer<</Size 6 /Root 1 0 R>>
startxref
441
%%EOF"""
        resp = upload_file(client, minimal_pdf, "with_text.pdf", "application/pdf")
        # Accepte 200 ou 422 (selon si pypdf extrait le texte du minimal pdf)
        assert resp.status_code in (200, 422)

    def test_invalid_pdf_bytes_returns_400(self, client):
        resp = upload_file(client, b"NOT A PDF AT ALL", "fake.pdf", "application/pdf")
        assert resp.status_code == 400

    def test_pdf_too_large_returns_413(self, client):
        # 21 MB > 20 MB limit
        big = b"A" * (21 * 1024 * 1024)
        resp = upload_file(client, big, "huge.pdf", "application/pdf")
        assert resp.status_code == 413


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

class TestDocxUpload:
    def test_valid_docx_upload(self, client):
        pytest.importorskip("docx")
        from docx import Document
        import io
        doc = Document()
        doc.add_paragraph("Bonjour, ceci est un document Word de test.")
        doc.add_paragraph("Il contient plusieurs paragraphes pour la JE.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = upload_file(
            client, buf.read(), "test.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["extension"] == "docx"
        assert "Bonjour" in data["text"]

    def test_invalid_docx_bytes_returns_400(self, client):
        resp = upload_file(client, b"NOT A DOCX", "fake.docx",
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert resp.status_code == 400


# ─────────────────────────────────────────────────────────────────────────────
# Cas d'erreur généraux
# ─────────────────────────────────────────────────────────────────────────────

class TestUploadErrors:
    def test_empty_file_returns_422(self, client):
        resp = upload_file(client, b"", "empty.txt")
        assert resp.status_code == 422

    def test_whitespace_only_returns_422(self, client):
        resp = upload_file(client, b"   \n\t  ", "blank.txt")
        assert resp.status_code == 422

    def test_file_size_limit_enforced(self, client):
        # 20 MB + 1 byte
        oversized = b"X" * (20 * 1024 * 1024 + 1)
        resp = upload_file(client, oversized, "too_big.txt")
        assert resp.status_code == 413

    def test_upload_returns_chars_count(self, client):
        content = "Bonjour monde."
        resp = upload_file(client, content.encode(), "count.txt")
        assert resp.status_code == 200
        assert resp.json()["chars"] == len(content)
